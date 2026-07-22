"""MCP tool handlers — each maps to a DRE capability.

These functions are called by the FastMCP server in server.py.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Optional


def render_document(
    markdown_content: str,
    template_name: str = "standard",
    output_path: Optional[str] = None,
    no_postprocess: bool = False,
) -> dict[str, Any]:
    """Convert markdown content to a DOCX file.

    Args:
        markdown_content: The full Markdown text to render.
        template_name: Name of the style template (without .yaml).
        output_path: Optional path for the output DOCX.
        no_postprocess: Skip OfficeCLI post-processing if True.

    Returns:
        dict with keys: success, output_path, errors, warnings
    """
    errors: list[str] = []
    warnings: list[str] = []

    try:
        from dre.parser.markdown_parser import MarkdownParser
        from dre.style.template import StyleTemplate
        from dre.renderer.docx_renderer import DocxRenderer

        # Determine template path
        tmpl_path = _template_path(template_name)
        if not tmpl_path.exists():
            return {
                "success": False,
                "output_path": "",
                "errors": [f"Template '{template_name}' not found"],
                "warnings": [],
            }

        # Parse
        parser = MarkdownParser()
        doc = parser.parse(markdown_content)

        # Load template
        template = StyleTemplate.from_yaml(tmpl_path)

        # Determine output path
        if not output_path:
            output_path = _default_output_name()

        # Render
        renderer = DocxRenderer(template)
        renderer.render(doc, output_path)
        output_path = str(Path(output_path).resolve())

        # Post-processing — explicit opt-in only
        if not no_postprocess:
            from dre.postprocess.officecli import OfficePostProcessor
            pp = OfficePostProcessor()
            if pp.is_available():
                try:
                    pp.refresh(output_path)
                except Exception:
                    warnings.append(
                        "OfficeCLI refresh failed. TOC page numbers will be correct "
                        "after opening the document in Word and pressing Ctrl+A → F9."
                    )
            else:
                # Not an error — TOC field is valid, just needs a refresh in Word
                pass

        return {
            "success": True,
            "output_path": output_path,
            "errors": errors,
            "warnings": warnings,
        }

    except Exception as exc:
        errors.append(str(exc))
        return {
            "success": False,
            "output_path": output_path or "",
            "errors": errors,
            "warnings": warnings,
        }


def validate_document(markdown_content: str) -> dict[str, Any]:
    """Parse and validate markdown content without rendering.

    Returns a summary of the document structure.
    """
    from dre.parser.markdown_parser import MarkdownParser

    errors: list[str] = []
    warnings: list[str] = []

    try:
        parser = MarkdownParser()
        doc = parser.parse(markdown_content)
        flat = doc.walk()

        summary: dict[str, int] = {}
        for node in flat:
            t = node.node_type
            summary[t] = summary.get(t, 0) + 1

        return {
            "valid": True,
            "title": doc.title or "",
            "node_count": len(flat),
            "summary": summary,
            "errors": errors,
            "warnings": warnings,
        }
    except Exception as exc:
        errors.append(str(exc))
        return {
            "valid": False,
            "title": "",
            "node_count": 0,
            "summary": {},
            "errors": errors,
            "warnings": warnings,
        }


def list_templates() -> dict[str, Any]:
    """List available style templates — both local and marketplace.

    Local templates are always returned.  Marketplace templates are fetched
    on a best-effort basis — if the network is unavailable only local
    results are returned with no error.

    Returns:
        dict with keys:
            local (list of dicts),
            marketplace (list of dicts, may be empty on network failure),
            all_names (combined sorted list for quick lookup)
    """
    from dre.config import templates_dir

    # ---- Local templates ------------------------------------------------
    td = templates_dir()
    yaml_files = sorted(td.glob("*.yaml")) if td.exists() else []
    local = []
    for f in yaml_files:
        try:
            from dre.style.template import StyleTemplate
            tmpl = StyleTemplate.from_yaml(f)
            meta = tmpl.get_metadata()
            local.append({
                "name": f.stem,
                "description": meta.get("description", ""),
                "file": str(f.resolve()),
                "source": "local",
            })
        except Exception:
            local.append({
                "name": f.stem,
                "description": "(failed to load)",
                "file": str(f.resolve()),
                "source": "local",
            })

    # ---- Marketplace (best-effort) --------------------------------------
    marketplace: list[dict] = []
    local_names = {t["name"] for t in local}
    try:
        import urllib.request
        import json
        url = "https://raw.githubusercontent.com/nihao-hello1/DRE-templates/main/index.json"
        req = urllib.request.Request(url)
        with urllib.request.urlopen(req, timeout=8) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        remote = data if isinstance(data, list) else data.get("templates", [])
        marketplace = [
            {
                "name": t["name"],
                "description": t.get("description", ""),
                "tags": t.get("tags", []),
                "version": t.get("version", "0.1.0"),
                "source": "marketplace",
            }
            for t in remote
            if t["name"] not in local_names  # don't duplicate local templates
        ]
    except Exception:
        pass  # Network unavailable — silent fallback, not an error

    # Combined sorted list for quick reference
    all_names = sorted(local_names | {t["name"] for t in marketplace})

    return {
        "local": local,
        "marketplace": marketplace,
        "marketplace_count": len(marketplace),
        "all_names": all_names,
        "total": len(local) + len(marketplace),
    }


def install_template(name: str) -> dict[str, Any]:
    """Install a template from the DRE marketplace.

    Downloads the YAML file to the local templates directory so it
    becomes available for rendering.  Call this when the user picks a
    marketplace template from the list_templates() results.

    Args:
        name: Template name (without .yaml).

    Returns:
        dict with keys: success, path, message/error
    """
    import urllib.request
    from dre.config import templates_dir

    local_path = templates_dir() / f"{name}.yaml"

    if local_path.exists():
        return {
            "success": True,
            "path": str(local_path),
            "message": f"Template '{name}' already installed.",
        }

    url = f"https://raw.githubusercontent.com/nihao-hello1/DRE-templates/main/templates/{name}.yaml"
    try:
        req = urllib.request.Request(url)
        with urllib.request.urlopen(req, timeout=15) as resp:
            content = resp.read()
        local_path.write_bytes(content)
        return {
            "success": True,
            "path": str(local_path),
            "message": f"Installed. Now available: render_document(..., template_name='{name}')",
        }
    except Exception as exc:
        return {
            "success": False,
            "path": "",
            "error": f"Download failed: {exc}",
        }


def get_document_info(docx_path: str) -> dict[str, Any]:
    """Return metadata about a rendered DOCX file."""
    from docx import Document as DocxDocument

    path = Path(docx_path)
    if not path.exists():
        return {"error": f"File not found: {docx_path}"}

    try:
        doc = DocxDocument(str(path))
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        section_count = len(doc.sections)
        char_count = sum(len(p.text) for p in doc.paragraphs)

        from docx.oxml.ns import qn
        headings = []
        for p in doc.paragraphs:
            pPr = p._p.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:outlineLvl")) is not None:
                headings.append({
                    "text": p.text[:100],
                    "level": int(pPr.find(qn("w:outlineLvl")).get(qn("w:val"))) + 1,
                })

        return {
            "file": str(path.resolve()),
            "file_size_bytes": path.stat().st_size,
            "paragraphs": para_count,
            "tables": table_count,
            "sections": section_count,
            "characters": char_count,
            "headings": headings[:20],
        }
    except Exception as exc:
        return {"error": str(exc)}


# ---------------------------------------------------------------------------
#  Internal helpers
# ---------------------------------------------------------------------------

def _template_path(template_name: str) -> Path:
    from dre.config import templates_dir
    # Backward compat: old name maps to standard
    name = "standard" if template_name == "tech_design" else template_name
    td = templates_dir()
    path = td / f"{name}.yaml"
    if not path.exists():
        raise FileNotFoundError(f"Template '{template_name}' not found in {td}")
    return path


def _default_output_name() -> str:
    import time
    return f"dre_output_{time.time_ns()}.docx"
