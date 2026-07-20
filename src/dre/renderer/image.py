"""Image rendering with captions."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from dre.ast.nodes import Image as ImageNode
from dre.style.defaults import ParagraphStyle

# Simple consecutive image counter
_image_counter: int = 0


def reset_image_counter() -> None:
    """Reset the image numbering counter (call at start of each render)."""
    global _image_counter
    _image_counter = 0


class ImageRenderer:
    """Renders Image AST nodes into a python-docx Document."""

    def __init__(self, docx: DocxDocument) -> None:
        self._docx = docx

    def render_image(
        self,
        node: ImageNode,
        caption_style: ParagraphStyle,
    ) -> None:
        """Embed an image and add its caption."""
        global _image_counter

        source_path = Path(node.source)
        if not source_path.exists():
            # Image not found — render placeholder text
            para = self._docx.add_paragraph()
            run = para.add_run(f"[图片: {node.alt_text}]")
            run.font.name = "SimSun"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.italic = True
            return

        # Add the image
        img_width = None
        if node.width:
            img_width = Cm(node.width / 36)  # approximate px to cm
        elif node.caption:
            img_width = Cm(14)  # default width

        run = self._docx.add_paragraph().add_run()
        pic = run.add_picture(
            str(source_path),
            width=img_width,
        )

        # Add caption below image
        if node.caption:
            _image_counter += 1
            caption_text = f"图{_image_counter} {node.caption}"

            para = self._docx.add_paragraph()
            para.paragraph_format.alignment = _align_from_str("center")
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(6)

            run = para.add_run(caption_text)
            run.font.name = caption_style.font_name
            run.font.size = _parse_pt(caption_style.font_size)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            _set_east_asian_font(run, caption_style.font_name)

    def render_image_inline(self, node: ImageNode, caption_style: ParagraphStyle) -> None:
        """Render an image that appeared inside a paragraph (inline markdown)."""
        self.render_image(node, caption_style)


def _parse_pt(value: str) -> Pt:
    val = value.strip().lower().replace("pt", "")
    return Pt(float(val))


def _set_east_asian_font(run, font_name: str) -> None:
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _align_from_str(align: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(align.strip().lower(), WD_ALIGN_PARAGRAPH.LEFT)
