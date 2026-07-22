"""Table-of-contents field-code insertion.

python-docx has no high-level TOC API, so we inject the raw OOXML
``w:fldChar`` / ``w:instrText`` elements via lxml.
"""

from __future__ import annotations

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from dre.style.defaults import ParagraphStyle, TOCConfig


def insert_toc(docx: DocxDocument, toc_config: TOCConfig, heading_style: ParagraphStyle) -> None:
    """Insert a TOC field at the end of the document (before first body content).

    Call this *after* the TOC title paragraph and *before* any body
    content — the TOC field code should be the first content element.
    """
    # Add a paragraph for the TOC field code
    toc_para = docx.add_paragraph()
    # Word TOC field instruction

    levels = toc_config.levels
    instr = f' TOC \\o "1-{levels}" \\h \\z \\u '

    # Build the field code using OOXML elements
    run = toc_para.add_run()

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")

    # Placeholder text that shows before Word updates the field
    placeholder = OxmlElement("w:t")
    placeholder.text = "（打开文档后按 Ctrl+A → F9 即可刷新目录页码）"
    placeholder.set(qn("xml:space"), "preserve")

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    run._r.append(placeholder)
    run._r.append(fldChar_end)

    # Move the TOC paragraph to the beginning of the document (right after the
    # TOC title which should already be there).  We do this by saving the body
    # element and re-inserting.
    body = docx.element.body
    # TOC paragraph should be the second element (after title paragraph)
    if len(body) > 1:
        toc_para_element = toc_para._element
        # Move it to position 1 (0 = title, 1 = TOC field)
        body.remove(toc_para_element)
        body.insert(1, toc_para_element)


def add_toc_title(docx: DocxDocument, toc_config: TOCConfig) -> None:
    """Add the TOC title paragraph at the very start of the document."""
    para = docx.add_paragraph()
    run = para.add_run(toc_config.title)
    run.font.name = toc_config.title_font_name
    run.font.size = _parse_pt(toc_config.title_font_size)
    run.bold = True
    _set_east_asian_font(run, toc_config.title_font_name)

    # Make sure it stays at position 0
    body = docx.element.body
    para_element = para._element
    body.remove(para_element)
    body.insert(0, para_element)
    return para


def _parse_pt(value: str) -> int:
    """Parse a pt string to python-docx Pt (int)."""
    val = value.strip().lower().replace("pt", "")
    return int(round(float(val) * 12700))


def _set_east_asian_font(run, font_name: str) -> None:
    """Set the East-Asian font on a run (for Chinese characters)."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
