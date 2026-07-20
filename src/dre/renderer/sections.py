"""Page-level document setup: margins, headers, footers, page numbers."""

from __future__ import annotations

import lxml.etree as ET

from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls

from dre.style.defaults import HeaderFooterContent, PageSetup


# ---------------------------------------------------------------------------
#  Page layout
# ---------------------------------------------------------------------------

def setup_page(docx: DocxDocument, page_setup: PageSetup) -> None:
    """Configure page size, orientation, and margins."""
    section = docx.sections[0]

    # Page size
    if page_setup.size.upper() == "A4":
        section.page_width = _mm_to_emu(210)
        section.page_height = _mm_to_emu(297)
    elif page_setup.size.upper() == "LETTER":
        section.page_width = _mm_to_emu(215.9)
        section.page_height = _mm_to_emu(279.4)
    # else: keep python-docx default

    # Orientation
    if page_setup.orientation and page_setup.orientation.lower() == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    # Margins
    section.top_margin = _parse_length(page_setup.margin_top)
    section.bottom_margin = _parse_length(page_setup.margin_bottom)
    section.left_margin = _parse_length(page_setup.margin_left)
    section.right_margin = _parse_length(page_setup.margin_right)


# ---------------------------------------------------------------------------
#  Headers & Footers
# ---------------------------------------------------------------------------

def setup_header(docx: DocxDocument, header: HeaderFooterContent) -> None:
    """Add header text to the default section."""
    if not header.text:
        return

    section = docx.sections[0]
    hdr = section.header
    hdr.is_linked_to_previous = False

    paragraph = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    paragraph.alignment = _alignment_from_str(header.alignment)

    run = paragraph.add_run(header.text)
    run.font.name = header.font_name
    run.font.size = _parse_pt(header.font_size)
    run.font.color.rgb = _rgb_from_hex("999999")  # subtle gray for headers
    # Enable East Asian font
    _set_east_asian_font(run, header.font_name)


def setup_footer(docx: DocxDocument, footer: HeaderFooterContent) -> None:
    """Add footer with optional page-number field codes."""
    section = docx.sections[0]
    ftr = section.footer
    ftr.is_linked_to_previous = False

    paragraph = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    paragraph.alignment = _alignment_from_str(footer.alignment)

    if footer.show_page_number:
        _insert_page_number_fields(paragraph, footer)
    elif footer.text:
        run = paragraph.add_run(footer.text)
        run.font.name = footer.font_name
        run.font.size = _parse_pt(footer.font_size)
        _set_east_asian_font(run, footer.font_name)


def _insert_page_number_fields(paragraph, footer: HeaderFooterContent) -> None:
    """Insert 'Page X of Y' fields using OOXML field codes.

    Handles formats like:
      "{page}"              → just page number
      "{page} / {numpages}" → page number / total pages
      "第 {page} 页"        → Chinese-style
    """
    fmt = footer.page_number_format

    # Split on markers
    parts = fmt.split("{page}")
    numparts = len(parts)
    numparts2_parts = [p.split("{numpages}") for p in parts] if "{numpages}" in fmt else None

    if numparts2_parts:
        # Contains both {page} and {numpages}
        for i, subparts in enumerate(numparts2_parts):
            if subparts[0]:
                run = paragraph.add_run(subparts[0])
                _style_footer_run(run, footer)
            if i < 2:  # after first split, before second (or after both)
                _add_field(paragraph, "PAGE", footer)
            if len(subparts) > 1 and i == 0:
                # {numpages} applies here
                _add_field(paragraph, "NUMPAGES", footer)
    elif "{page}" in fmt:
        for i, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                _style_footer_run(run, footer)
            if i < len(parts) - 1:
                _add_field(paragraph, "PAGE", footer)
    else:
        # Static text
        run = paragraph.add_run(fmt)
        _style_footer_run(run, footer)


def _add_field(paragraph, field_name: str, footer: HeaderFooterContent) -> None:
    """Insert a Word field code (PAGE or NUMPAGES) as a run."""
    run = paragraph.add_run()
    _style_footer_run(run, footer)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = field_name

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    # Word inserts the computed value between separate and end —
    # we leave it empty; OfficeCLI or Ctrl+A F9 will fill it.
    run._r.append(fldChar_end)


def _style_footer_run(run, footer: HeaderFooterContent) -> None:
    run.font.name = footer.font_name
    run.font.size = _parse_pt(footer.font_size)
    run.font.color.rgb = _rgb_from_hex("999999")
    _set_east_asian_font(run, footer.font_name)


# ---------------------------------------------------------------------------
#  Formatting helpers
# ---------------------------------------------------------------------------

def _mm_to_emu(mm: float) -> int:
    """Convert millimetres to EMU (English Metric Units, used by OOXML)."""
    return int(mm * 36000)


def _cm_to_emu(cm: float) -> int:
    """Convert centimetres to EMU."""
    return int(cm * 360000)


def _parse_length(value: str) -> int:
    """Parse a CSS-style length string into EMU.

    Supported units: cm, mm, pt, in.  Falls back to cm if no unit.
    """
    value = value.strip().lower()
    if value.endswith("cm"):
        return _cm_to_emu(float(value.replace("cm", "")))
    elif value.endswith("mm"):
        return _mm_to_emu(float(value.replace("mm", "")))
    elif value.endswith("pt"):
        return int(float(value.replace("pt", "")) * 12700)
    elif value.endswith("in"):
        return int(float(value.replace("in", "")) * 914400)
    else:
        return _cm_to_emu(float(value))


def _parse_pt(value: str) -> int:
    """Parse a pt string to python-docx Pt (int)."""
    val = value.strip().lower().replace("pt", "")
    return int(round(float(val) * 12700))


def _alignment_from_str(align: str):
    """Map alignment string to python-docx WD_ALIGN_PARAGRAPH."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(align.strip().lower(), WD_ALIGN_PARAGRAPH.LEFT)


def _rgb_from_hex(hex_color: str):
    """Convert hex string like '1F4E79' to a python-docx RGBColor."""
    from docx.shared import RGBColor
    hex_color = hex_color.lstrip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _set_east_asian_font(run, font_name: str) -> None:
    """Set the East-Asian font on a run (for Chinese characters)."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
