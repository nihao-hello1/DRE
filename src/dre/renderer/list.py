"""List rendering (bulleted and ordered)."""

from __future__ import annotations

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from dre.ast.nodes import BulletList, DocumentNode, OrderedList
from dre.style.defaults import ParagraphStyle


class ListRenderer:
    """Renders BulletList and OrderedList AST nodes."""

    def __init__(self, docx: DocxDocument) -> None:
        self._docx = docx
        self._list_counters: dict[int, int] = {}

    def render_bullet_list(self, node: BulletList, style: ParagraphStyle, level: int = 0) -> None:
        """Render a bullet list, supporting nesting."""
        for item_children in node.items:
            if not item_children:
                continue

            para = self._docx.add_paragraph()
            self._apply_list_format(para, style, level)

            # Bullet character
            bullet_chars = ["●", "○", "■"][level % 3]
            run = para.add_run(f"{bullet_chars}  ")
            run.font.name = style.font_name
            run.font.size = _parse_pt(style.font_size)
            _set_east_asian_font(run, style.font_name)

            # Render item content
            for child in item_children:
                if hasattr(child, "text") and child.text:
                    run = para.add_run(child.text)
                    run.font.name = style.font_name
                    run.font.size = _parse_pt(style.font_size)
                    _set_east_asian_font(run, style.font_name)

                # Nested lists
                if isinstance(child, (BulletList, OrderedList)):
                    # Remove the outer paragraph's bullet text for this child
                    if isinstance(child, BulletList):
                        self.render_bullet_list(child, style, level + 1)
                    else:
                        self.render_ordered_list(child, style, level + 1)

    def render_ordered_list(self, node: OrderedList, style: ParagraphStyle, level: int = 0) -> None:
        """Render an ordered list, supporting nesting."""
        counter = node.start
        for item_children in node.items:
            if not item_children:
                continue

            para = self._docx.add_paragraph()
            self._apply_list_format(para, style, level)

            # Number prefix
            number_formats = [
                f"{counter}.",
                f"{chr(96 + counter)}.",
                f"{counter}.",
            ]
            prefix = number_formats[level % 3]
            run = para.add_run(f"{prefix}  ")
            run.font.name = style.font_name
            run.font.size = _parse_pt(style.font_size)
            _set_east_asian_font(run, style.font_name)

            for child in item_children:
                if hasattr(child, "text") and child.text:
                    run = para.add_run(child.text)
                    run.font.name = style.font_name
                    run.font.size = _parse_pt(style.font_size)
                    _set_east_asian_font(run, style.font_name)

                if isinstance(child, (BulletList, OrderedList)):
                    if isinstance(child, BulletList):
                        self.render_bullet_list(child, style, level + 1)
                    else:
                        self.render_ordered_list(child, style, level + 1)

            counter += 1

    def _apply_list_format(self, para, style: ParagraphStyle, level: int) -> None:
        """Set paragraph formatting for a list item."""
        pf = para.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = _parse_length(style.space_before)
        pf.space_after = _parse_length(style.space_after)
        pf.line_spacing = style.line_spacing

        # Indent based on nesting level
        indent = 0.5 + (level * 0.8)
        pf.left_indent = Pt(indent * 12)
        pf.first_line_indent = Pt(-0.3 * 12)


def _parse_pt(value: str) -> Pt:
    val = value.strip().lower().replace("pt", "")
    return Pt(float(val))


def _parse_length(value):
    from docx.shared import Pt as DxPt
    value = value.strip().lower()
    if value.endswith("pt"):
        return DxPt(float(value.replace("pt", "")))
    return DxPt(float(value.replace("cm", "")) * 28.35)


def _set_east_asian_font(run, font_name: str) -> None:
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
