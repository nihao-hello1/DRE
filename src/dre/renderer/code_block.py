"""Code block rendering with shaded background."""

from __future__ import annotations

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from dre.ast.nodes import CodeBlock
from dre.style.defaults import ParagraphStyle


class CodeBlockRenderer:
    """Renders CodeBlock AST nodes with monospaced font and background."""

    def __init__(self, docx: DocxDocument) -> None:
        self._docx = docx

    def render_code_block(self, node: CodeBlock, style: ParagraphStyle) -> None:
        """Add a formatted code block to the document."""
        if not node.code:
            return

        # Language label
        if node.language:
            lang_para = self._docx.add_paragraph()
            lang_para.paragraph_format.space_before = Pt(6)
            lang_para.paragraph_format.space_after = Pt(0)
            lang_para.paragraph_format.left_indent = Pt(6)
            run = lang_para.add_run(f"  {node.language}")
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.italic = True

        # Split code into lines and render each as a separate paragraph
        lines = node.code.split("\n")
        for line_idx, line in enumerate(lines):
            para = self._docx.add_paragraph()
            self._apply_code_format(para, style, is_last=(line_idx == len(lines) - 1))

            run = para.add_run(line if line else " ")
            run.font.name = style.font_name
            run.font.size = _parse_pt(style.font_size)
            run.font.color.rgb = _parse_color(style.color)

            # Background on the run
            _set_run_shading(run, "F5F5F5")

    def _apply_code_format(self, para, style: ParagraphStyle, is_last: bool = False) -> None:
        """Set paragraph formatting for code lines."""
        pf = para.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = style.line_spacing
        pf.left_indent = _parse_length(style.space_before) if style.space_before != "0pt" else Pt(12)
        pf.right_indent = Pt(6)

        # Top border on first line, bottom on last
        if not hasattr(self, "_first_code_line"):
            self._first_code_line = True
        else:
            self._first_code_line = False


def _set_run_shading(run, hex_color: str) -> None:
    """Apply background shading to a run."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def _parse_pt(value: str) -> Pt:
    val = value.strip().lower().replace("pt", "")
    return Pt(float(val))


def _parse_color(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


def _parse_length(value):
    from docx.shared import Pt as DxPt
    value = value.strip().lower()
    if value.endswith("pt"):
        return DxPt(float(value.replace("pt", "")))
    return DxPt(float(value.replace("cm", "")) * 28.35)
