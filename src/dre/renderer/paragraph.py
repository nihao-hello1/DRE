"""Paragraph and heading rendering."""

from __future__ import annotations

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from dre.ast.nodes import Heading, Paragraph, TextRun
from dre.style.defaults import ParagraphStyle


class ParagraphRenderer:
    """Renders Paragraph and Heading AST nodes into a python-docx Document."""

    def __init__(self, docx: DocxDocument) -> None:
        self._docx = docx

    def render_heading(self, node: Heading, style: ParagraphStyle) -> None:
        """Render a heading with the resolved style."""
        para = self._docx.add_paragraph()
        self._apply_paragraph_format(para, style)

        run = para.add_run(node.text)
        self._apply_run_format(run, style)

        # Outline level for TOC — set via raw OOXML
        if style.outline_level is not None:
            pPr = para._p.get_or_add_pPr()
            ol = pPr.find(qn("w:outlineLvl"))
            if ol is None:
                ol = OxmlElement("w:outlineLvl")
                pPr.append(ol)
            ol.set(qn("w:val"), str(style.outline_level - 1))

    def render_paragraph(self, node: Paragraph, style: ParagraphStyle) -> None:
        """Render a body paragraph, including inline formatting."""
        if node.is_empty():
            return

        para = self._docx.add_paragraph()
        self._apply_paragraph_format(para, style)

        if node.children:
            for text_run_node in node.children:
                self._render_text_run(para, text_run_node, style)
        else:
            run = para.add_run(node.text)
            self._apply_run_format(run, style)

    def render_empty_line(self) -> None:
        """Add a minimal spacer paragraph."""
        para = self._docx.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

    # ------------------------------------------------------------------
    #  Internal helpers
    # ------------------------------------------------------------------

    def _render_text_run(self, para, node: TextRun, parent_style: ParagraphStyle) -> None:
        """Add a formatted TextRun to a paragraph."""
        if not node.text:
            return

        run = para.add_run(node.text)
        run.bold = node.bold
        run.italic = node.italic
        run.font.name = parent_style.font_name

        if node.code:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            # Light gray background via shading
            rPr = run._r.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F0F0F0")
            rPr.append(shd)
        else:
            run.font.size = _parse_pt(parent_style.font_size)
            run.font.color.rgb = _parse_color(parent_style.color)

        self._set_east_asian_font(run, parent_style.font_name)

        if node.link:
            # Add hyperlink using OOXML
            self._make_hyperlink(para, node.link, node.text, parent_style)

    def _apply_paragraph_format(self, para, style: ParagraphStyle) -> None:
        """Set paragraph-level formatting from a ParagraphStyle."""
        pf = para.paragraph_format

        pf.alignment = _align_from_str(style.alignment)
        pf.space_before = _parse_length(style.space_before)
        pf.space_after = _parse_length(style.space_after)
        pf.line_spacing = style.line_spacing

        if style.first_line_indent and style.first_line_indent != "0cm":
            pf.first_line_indent = _parse_length(style.first_line_indent)

        if style.left_indent and style.left_indent != "0cm":
            pf.left_indent = _parse_length(style.left_indent)

        if style.keep_with_next:
            pf.keep_with_next = True

        if style.page_break_before:
            pf.page_break_before = True

    def _apply_run_format(self, run, style: ParagraphStyle) -> None:
        """Set run-level formatting from a ParagraphStyle."""
        run.font.name = style.font_name
        run.font.size = _parse_pt(style.font_size)
        run.bold = style.bold
        run.italic = style.italic

        if style.color:
            run.font.color.rgb = _parse_color(style.color)

        self._set_east_asian_font(run, style.font_name)

    def _set_east_asian_font(self, run, font_name: str) -> None:
        """Set the East-Asian font for CJK character rendering."""
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)

    def _make_hyperlink(self, paragraph, url: str, text: str, style: ParagraphStyle) -> None:
        """Add an OOXML hyperlink to the paragraph."""
        from docx.oxml import OxmlElement

        # Create the hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), "")
        hyperlink.set(qn("w:history"), "1")

        # Create a new run inside the hyperlink
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rPr.append(OxmlElement("w:u"))
        # Set the /last character
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(rPr)
        r.append(t)
        hyperlink.append(r)

        # Append hyperlink to paragraph
        paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
#  Unit conversion helpers
# ---------------------------------------------------------------------------

def _parse_pt(value: str) -> Pt:
    """Convert '12pt' → Pt(12)."""
    val = value.strip().lower().replace("pt", "")
    return Pt(float(val))


def _parse_length(value: str):
    """Convert CSS-style length to python-docx Cm/Pt/Inches."""
    from docx.shared import Cm, Pt as DxPt, Inches

    value = value.strip().lower()
    if value.endswith("cm"):
        return Cm(float(value.replace("cm", "")))
    elif value.endswith("mm"):
        return Cm(float(value.replace("mm", "")) / 10)
    elif value.endswith("pt"):
        return DxPt(float(value.replace("pt", "")))
    elif value.endswith("in"):
        return Inches(float(value.replace("in", "")))
    # Assume cm
    return Cm(float(value))


def _align_from_str(align: str):
    """Map alignment name to WD_ALIGN_PARAGRAPH."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(align.strip().lower(), WD_ALIGN_PARAGRAPH.LEFT)


def _parse_color(hex_color: str) -> RGBColor:
    """Convert '1F4E79' → RGBColor(...)."""
    hex_color = hex_color.lstrip("#")
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )
