"""Table rendering."""

from __future__ import annotations

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

from dre.ast.nodes import Table as TableNode
from dre.style.defaults import ParagraphStyle, TableStyle
from dre.renderer.paragraph import ParagraphRenderer


class TableRenderer:
    """Renders Table AST nodes into a python-docx Document."""

    def __init__(self, docx: DocxDocument) -> None:
        self._docx = docx

    def render_table(
        self,
        node: TableNode,
        style: TableStyle,
        caption_style: ParagraphStyle,
        body_style: ParagraphStyle,
    ) -> None:
        """Add a formatted table to the document."""
        if not node.headers and not node.rows:
            return

        # Caption before table
        if node.caption:
            pr = ParagraphRenderer(self._docx)
            # Use a simple paragraph with caption style
            para = self._docx.add_paragraph()
            run = para.add_run(node.caption)
            run.font.name = caption_style.font_name
            run.font.size = _parse_pt(caption_style.font_size)
            run.bold = False
            run.font.color.rgb = _parse_color("333333")
            _set_east_asian_font(run, caption_style.font_name)
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.alignment = _align_from_str("left")

        # Determine column count
        col_count = len(node.headers) if node.headers else max((len(r) for r in node.rows), default=1)

        # Create table
        table = self._docx.add_table(
            rows=len(node.rows) + (1 if node.headers else 0),
            cols=col_count,
        )

        # Apply a built-in table style as base
        try:
            table.style = self._docx.styles["Light Grid Accent 1"]
        except KeyError:
            pass  # Fall back to default table style

        # Render header row
        if node.headers:
            for i, header_text in enumerate(node.headers):
                cell = table.rows[0].cells[i]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(header_text)

                # Format header
                run.font.name = style.font_name
                run.font.size = _parse_pt(style.font_size)
                run.bold = style.header_bold
                run.font.color.rgb = _parse_color(style.header_font_color)
                _set_east_asian_font(run, style.font_name)
                para.paragraph_format.alignment = _align_from_str("center")
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)

                # Header background color
                _set_cell_shading(cell, style.header_bg)

        # Render data rows
        row_offset = 1 if node.headers else 0
        for row_idx, row_data in enumerate(node.rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx >= col_count:
                    continue
                cell = table.rows[row_idx + row_offset].cells[col_idx]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(cell_text)

                run.font.name = style.font_name
                run.font.size = _parse_pt(style.font_size)
                _set_east_asian_font(run, style.font_name)
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)

                # Alternate row shading
                if row_idx % 2 == 1:
                    _set_cell_shading(cell, style.row_alt_color)

        # Table borders
        _set_table_borders(table, style.border_color)

        # Spacing after table
        last_row = table.rows[-1]
        for cell in last_row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
#  OOXML helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_color: str) -> None:
    """Set cell background shading color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def _set_table_borders(table, hex_color: str) -> None:
    """Set table borders to a single color."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), hex_color)
        tblBorders.append(border)

    tblPr.append(tblBorders)


def _parse_pt(value: str) -> Pt:
    val = value.strip().lower().replace("pt", "")
    return Pt(float(val))


def _parse_color(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


def _set_east_asian_font(run, font_name: str) -> None:
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _align_from_str(align: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(align.strip().lower(), WD_ALIGN_PARAGRAPH.LEFT)
