"""Core pipeline tests — run with: pytest tests/"""

import os
import tempfile
import pytest
from docx import Document
from docx.oxml.ns import qn

from dre.mcp_server.tools import (
    list_templates,
    validate_document,
    render_document,
    get_document_info,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def simple_md():
    return "# Test Title\n\n## Section One\n\nBody text here.\n\n### Subsection\n\nMore text."


@pytest.fixture
def full_md():
    return (
        "# 测试方案\n\n"
        "## 项目概述\n\n正文内容。\n\n"
        "## 技术方案\n\n"
        "### 架构\n\n内容。\n\n"
        "| 设备 | 型号 | 数量 |\n"
        "|------|------|------|\n"
        "| Spine | S6860 | 2 |\n\n"
        "```bash\nrouter ospf 100\n```\n\n"
        "> 注意：割接窗口建议凌晨。\n\n"
    )


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestTemplates:
    def test_list_has_standard(self):
        r = list_templates()
        names = {t["name"] for t in r["templates"]}
        assert "standard" in names
        assert "formal" in names
        assert "compact" in names
        assert "modern" in names
        assert "mac_standard" in names

    def test_tech_design_fallback(self, simple_md):
        r = render_document(simple_md, template_name="tech_design", no_postprocess=True)
        assert r["success"]
        os.remove(r["output_path"])


class TestValidation:
    def test_basic(self, simple_md):
        r = validate_document(simple_md)
        assert r["valid"]
        assert r["title"] == "Test Title"
        assert r["node_count"] > 0

    def test_empty(self):
        r = validate_document("")
        assert r["valid"]


class TestRendering:
    def test_render_creates_docx(self, simple_md):
        r = render_document(simple_md, no_postprocess=True)
        assert r["success"]
        assert r["output_path"].endswith(".docx")
        assert os.path.exists(r["output_path"])
        os.remove(r["output_path"])

    def test_page_margins(self, simple_md):
        r = render_document(simple_md, template_name="standard", no_postprocess=True)
        doc = Document(r["output_path"])
        s = doc.sections[0]
        assert abs(s.top_margin / 360000 - 2.54) < 0.05
        os.remove(r["output_path"])

    def test_compact_margins_are_smaller(self, simple_md):
        import time
        r1 = render_document(simple_md, template_name="compact",
                             output_path=f"test_compact_{time.time_ns()}.docx", no_postprocess=True)
        r2 = render_document(simple_md, template_name="formal",
                             output_path=f"test_formal_{time.time_ns()}.docx", no_postprocess=True)
        d1 = Document(r1["output_path"]).sections[0]
        d2 = Document(r2["output_path"]).sections[0]
        assert abs(d1.top_margin / 360000 - 2.0) < 0.05, f"compact top={d1.top_margin/360000:.2f}"
        assert abs(d2.top_margin / 360000 - 3.2) < 0.05, f"formal top={d2.top_margin/360000:.2f}"
        os.remove(r1["output_path"])
        os.remove(r2["output_path"])

    def test_headings_have_numbering(self, simple_md):
        r = render_document(simple_md, no_postprocess=True)
        doc = Document(r["output_path"])
        headings = 0
        for p in doc.paragraphs:
            pPr = p._p.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:outlineLvl")) is not None:
                headings += 1
                np = pPr.find(qn("w:numPr"))
                assert np is not None, f"Heading '{p.text[:20]}' missing numPr"
        assert headings >= 1
        os.remove(r["output_path"])

    def test_numbering_ilvl_h2_is_0(self, full_md):
        """When H1 is doc title, H2 should use ilvl=0 (not ilvl=1)."""
        r = render_document(full_md, no_postprocess=True)
        doc = Document(r["output_path"])
        for p in doc.paragraphs:
            pPr = p._p.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:outlineLvl")) is not None:
                ol = pPr.find(qn("w:outlineLvl"))
                np = pPr.find(qn("w:numPr"))
                ilvl = int(np.find(qn("w:ilvl")).get(qn("w:val")))
                lvl = int(ol.get(qn("w:val")))
                if lvl == 0:  # outlineLvl=0 means H1 in Word, but H2 in our numbering
                    # When H1 is title, H2 should have ilvl=0
                    assert ilvl == 0, f"H2 ilvl should be 0, got {ilvl}"
        os.remove(r["output_path"])

    def test_body_formatting(self, simple_md):
        r = render_document(simple_md, template_name="standard", no_postprocess=True)
        doc = Document(r["output_path"])
        for p in doc.paragraphs:
            pPr = p._p.find(qn("w:pPr"))
            if pPr is None or pPr.find(qn("w:outlineLvl")) is not None:
                continue
            if p.text.strip() and len(p.text) > 5:
                rn = p.runs[0]
                assert rn.font.name == "SimSun", f"Body font should be SimSun"
                assert rn.font.size == 152400  # 12pt * 12700
                break
        os.remove(r["output_path"])

    def test_all_four_templates_render(self, full_md):
        for t in ["standard", "formal", "compact", "modern"]:
            r = render_document(full_md, template_name=t, no_postprocess=True)
            assert r["success"], f"Template {t} failed"
            assert os.path.exists(r["output_path"])
            os.remove(r["output_path"])


class TestDocumentInfo:
    def test_returns_headings(self, simple_md):
        r = render_document(simple_md, no_postprocess=True)
        info = get_document_info(r["output_path"])
        assert len(info.get("headings", [])) > 0
        assert info["paragraphs"] > 0
        os.remove(r["output_path"])
